"""Microsoft Word (.docx) loader using python-docx."""
from __future__ import annotations
import io
import uuid
from pathlib import Path

from cognity_ai.loaders.base import BaseLoader
from cognity_ai.models.document import Document, ImageRef

# Heading styles recognized for building page_map sections
_HEADING_STYLES = {"heading 1", "heading 2", "heading 3", "title"}


class DocxLoader(BaseLoader):
    """Loads .docx files. Heading paragraphs become page_map entries."""

    @property
    def supported_extensions(self) -> list[str]:
        return [".docx", ".doc"]

    def load(self, path: str) -> list[Document]:
        try:
            import docx  # type: ignore
        except ImportError:
            raise ImportError(
                "python-docx is required to load .docx files. "
                "Install it with: pip install python-docx"
            )

        p = Path(path)
        doc = docx.Document(str(p))
        doc_id = p.stem + "_" + uuid.uuid4().hex[:8]

        text_parts: list[str] = []
        page_map: list[dict] = []
        image_refs: list[ImageRef] = []
        char_offset = 0
        section_num = 0

        current_heading = ""
        section_start = 0

        def _close_section(end: int) -> None:
            nonlocal section_num
            if end > section_start:
                page_map.append({
                    "page_num": section_num,
                    "start_char": section_start,
                    "end_char": end,
                    "heading": current_heading,
                })
                section_num += 1

        for para in doc.paragraphs:
            style_name = (para.style.name or "").lower()
            para_text = para.text

            if style_name in _HEADING_STYLES:
                # Close previous section
                _close_section(char_offset)
                current_heading = para_text
                section_start = char_offset

            line = para_text + "\n"
            text_parts.append(line)
            char_offset += len(line)

        # Close final section
        _close_section(char_offset)

        if not page_map:
            page_map = [{"page_num": 0, "start_char": 0, "end_char": char_offset, "heading": ""}]

        # Extract inline images
        image_index = 0
        try:
            for shape in doc.inline_shapes:
                try:
                    rId = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
                    image_part = doc.part.related_parts[rId]
                    img_bytes = image_part.blob
                    mime = image_part.content_type or "image/png"
                    image_refs.append(ImageRef(
                        image_id=f"{p.stem}_img{image_index}",
                        char_offset=0,
                        image_bytes=img_bytes,
                        mime_type=mime,
                        page_num=0,
                    ))
                    image_index += 1
                except Exception:
                    pass
        except Exception:
            pass

        full_text = "".join(text_parts)

        return [
            Document(
                doc_id=doc_id,
                text=full_text,
                source_path=str(p.resolve()),
                source_name=p.name,
                loader="DocxLoader",
                file_extension=p.suffix.lower(),
                file_size_bytes=p.stat().st_size,
                page_count=len(page_map),
                page_map=page_map,
                image_refs=image_refs,
            )
        ]
